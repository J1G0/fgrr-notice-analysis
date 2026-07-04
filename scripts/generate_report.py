"""전처리 분석 보고서 생성"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from dotenv import load_dotenv
import psycopg2

load_dotenv()

conn = psycopg2.connect(os.environ["DATABASE_URL"])

def q(sql):
    with conn.cursor() as cur:
        cur.execute(sql)
        return cur.fetchall()

def q1(sql):
    r = q(sql)
    return r[0][0] if r else None


# ── 데이터 수집 ──────────────────────────────────────────────────
total_rows = q1("SELECT COUNT(*) FROM fgrr_notices")
total_files = q1('SELECT COUNT(DISTINCT "파일명") FROM fgrr_notices')

columns = q("""
    SELECT column_name, data_type
    FROM information_schema.columns
    WHERE table_name='fgrr_notices' AND table_schema='public'
    ORDER BY ordinal_position
""")
col_names = [c[0] for c in columns if c[0] != 'id']

null_stats = []
for col in col_names:
    not_null = q1(f'SELECT COUNT(*) FROM fgrr_notices WHERE "{col}" IS NOT NULL')
    null_pct = round((1 - not_null / total_rows) * 100, 1) if total_rows else 0
    null_stats.append((col, not_null, null_pct))

유형_dist = q("""
    SELECT "유형", COUNT(*) FROM fgrr_notices
    GROUP BY "유형" ORDER BY COUNT(*) DESC LIMIT 10
""")

기관_dist = q("""
    SELECT "지방산림청", COUNT(*) FROM fgrr_notices
    GROUP BY "지방산림청" ORDER BY COUNT(*) DESC
""")

지정유형_dist = q("""
    SELECT "지정유형", COUNT(*) FROM fgrr_notices
    WHERE "지정유형" IS NOT NULL
    GROUP BY "지정유형" ORDER BY COUNT(*) DESC LIMIT 10
""")

유전자원여부 = q("""
    SELECT "유전자원보호구역_여부", COUNT(*) FROM fgrr_notices
    GROUP BY "유전자원보호구역_여부"
""")

날짜패턴 = q("""
    SELECT
      CASE
        WHEN "고시날짜" ~ '^\d{4}년 \d{2}월 \d{2}일$' THEN 'YYYY년 MM월 DD일 (제로패딩)'
        WHEN "고시날짜" ~ '^\d{4}년 \d{1,2}월 \d{1,2}일$' THEN 'YYYY년 M월 D일 (비패딩)'
        WHEN "고시날짜" ~ '^\d{4}-\d{2}-\d{2}$' THEN 'YYYY-MM-DD'
        WHEN "고시날짜" ~ '^\d{4}\. \d{1,2}\. \d{1,2}\.$' THEN 'YYYY. M. D.'
        WHEN "고시날짜" ~ '^\d{4}년\d{2}월\d{2}일$' THEN 'YYYY년MM월DD일 (공백없음)'
        WHEN "고시날짜" ~ '^\d{4}\.\d{2}\.\d{2}' THEN 'YYYY.MM.DD'
        ELSE '기타'
      END AS 패턴,
      COUNT(*)
    FROM fgrr_notices
    WHERE "고시날짜" IS NOT NULL
    GROUP BY 1 ORDER BY 2 DESC
""")

면적단위 = q("""
    SELECT
      CASE
        WHEN "지정면적" ~ '[㎡]' THEN '㎡ 명시'
        WHEN "지정면적" ~ 'ha' THEN 'ha 명시'
        WHEN "지정면적" ~ '^[\d,]+$' THEN '순수 숫자'
        WHEN "지정면적" IS NULL THEN 'NULL'
        ELSE '기타'
      END AS 유형,
      COUNT(*)
    FROM fgrr_notices
    GROUP BY 1 ORDER BY 2 DESC
""")

소재지패턴 = q("""
    SELECT
      CASE
        WHEN "소재지" ~ '^(서울|부산|대구|인천|광주|대전|울산|세종|경기|강원|충북|충남|전북|전남|경북|경남|제주)'
          THEN '시도 포함'
        WHEN "소재지" IS NULL THEN 'NULL'
        ELSE '시도 미포함'
      END AS 패턴,
      COUNT(*)
    FROM fgrr_notices
    GROUP BY 1 ORDER BY 2 DESC
""")

지번패턴 = q("""
    SELECT
      CASE
        WHEN "지번" ~ '^산\d+' THEN '산+숫자'
        WHEN "지번" ~ '^산 \d+' THEN '산+공백+숫자'
        WHEN "지번" ~ '^\d+' THEN '숫자시작'
        WHEN "지번" IS NULL THEN 'NULL'
        ELSE '기타'
      END AS 패턴,
      COUNT(*)
    FROM fgrr_notices
    GROUP BY 1 ORDER BY 2 DESC
""")

raw_text_stats = q("""
    SELECT
      MIN(LENGTH("raw_text")),
      MAX(LENGTH("raw_text")),
      ROUND(AVG(LENGTH("raw_text"))),
      PERCENTILE_CONT(0.5) WITHIN GROUP (ORDER BY LENGTH("raw_text"))
    FROM fgrr_notices WHERE "raw_text" IS NOT NULL
""")[0]

동적컬럼 = [
    '구분', '시도', '시군', '읍면', '리동', '위치',
    '산림유전자원보호림번호', '보호대상수종', '지정기간', '관리소',
    '연번', '세부유형', '정정사유', '지정기간'
]
동적_stats = []
for col in 동적컬럼:
    try:
        cnt = q1(f'SELECT COUNT(*) FROM fgrr_notices WHERE "{col}" IS NOT NULL')
        pct = round(cnt / total_rows * 100, 1)
        동적_stats.append((col, cnt, pct))
    except Exception:
        conn.rollback()

유형_unique = q1('SELECT COUNT(DISTINCT "유형") FROM fgrr_notices')
지정유형_unique = q1('SELECT COUNT(DISTINCT "지정유형") FROM fgrr_notices WHERE "지정유형" IS NOT NULL')

conn.close()


# ── 문서 작성 ──────────────────────────────────────────────────
doc = Document()

# 여백 설정
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

def set_font(run, size=10, bold=False, color=None):
    run.font.name = '맑은 고딕'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = '맑은 고딕'
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D) if level == 1 else RGBColor(0x2E, 0x75, 0xB6)
    return p

def para(doc, text, size=10):
    p = doc.add_paragraph(text)
    for run in p.runs:
        set_font(run, size)
    return p

def table_with_header(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # 헤더
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_font(run, 9, bold=True, color=(0x1F, 0x49, 0x7D))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell._tc.get_or_add_tcPr()
    # 데이터
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val) if val is not None else ''
            for run in cell.paragraphs[0].runs:
                set_font(run, 9)
    return t


# ── 표지 ──
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('유전자원보호구역 관보고시 DB 구축')
set_font(run, 18, bold=True, color=(0x1F, 0x49, 0x7D))

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run('전처리 품질 분석 보고서')
set_font(run, 14, color=(0x44, 0x72, 0xC4))

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('2026년 6월')
set_font(run, 11)

doc.add_page_break()


# ── 1. 개요 ──
heading(doc, '1. 개요')
para(doc, f"""산림청 산하 5개 지방산림청의 산림유전자원보호구역 관련 관보고시 PDF {total_files}건을 처리하여 \
PostgreSQL DB에 적재하였다. 총 {total_rows:,}개 필지 행이 생성되었으며, Claude API를 활용한 비정형 \
PDF 파싱 방식으로 인해 고시별 컬럼 구조의 차이가 존재한다. 본 보고서는 적재 결과의 품질을 분석하고 \
후속 전처리 방향을 제시한다.""")

doc.add_paragraph()
table_with_header(doc,
    ['항목', '값'],
    [
        ['처리 PDF 수', f'{total_files}건'],
        ['총 필지 수', f'{total_rows:,}행'],
        ['테이블 컬럼 수', f'{len(col_names)}개'],
        ['기본 설계 컬럼 수', '19개'],
        ['동적 추가 컬럼 수', f'{len(col_names) - 19}개'],
    ]
)


# ── 2. 컬럼 현황 ──
doc.add_page_break()
heading(doc, '2. 컬럼 현황')

heading(doc, '2.1 기본 컬럼 null 비율', level=2)
para(doc, '19개 기본 컬럼의 데이터 충전율이다. 고시 유형에 따라 존재하지 않는 필드(해제사유, 잔여면적 등)는 \
구조적 null이므로 데이터 손실과 구분이 필요하다.')
doc.add_paragraph()

base_cols = ['번호', '파일명', '지방산림청', '고시번호', '유형', '고시날짜',
             '소재지', '지번', '임반', '지목', '소유자', '지적',
             '지정면적', '해제면적', '잔여면적', '지정유형', '해제사유', '명칭', '비고']
base_stat_rows = [(col, nn, f'{pct}%') for col, nn, pct in null_stats if col in base_cols]
table_with_header(doc,
    ['컬럼명', '유효값 수', 'null 비율'],
    base_stat_rows
)

doc.add_paragraph()
heading(doc, '2.2 동적 추가 컬럼 현황', level=2)
para(doc, '고시별 표 구조 차이로 인해 파싱 중 동적으로 추가된 컬럼들이다. \
대부분 특정 소수 고시에서만 나타나는 필드로 null 비율이 매우 높다.')
doc.add_paragraph()
table_with_header(doc,
    ['컬럼명', '유효값 수', '충전율'],
    [(col, cnt, f'{pct}%') for col, cnt, pct in 동적_stats]
)


# ── 3. 주요 컬럼 품질 분석 ──
doc.add_page_break()
heading(doc, '3. 주요 컬럼 품질 분석')

heading(doc, '3.1 유형 컬럼', level=2)
para(doc, f'현재 유형 컬럼에 {유형_unique}개의 distinct 값이 존재한다. 동일한 고시 유형이 \
파일명 표기 방식, 띄어쓰기, 괄호 종류 차이 등으로 다양하게 기록되어 정규화가 필요하다.')
doc.add_paragraph()
table_with_header(doc,
    ['유형 값 (상위 10)', '건수'],
    [(v, c) for v, c in 유형_dist]
)

doc.add_paragraph()
heading(doc, '3.2 지정유형 컬럼', level=2)
para(doc, f'필지별 보호구역 유형을 나타내는 핵심 컬럼이다. {지정유형_unique}개 distinct 값이 존재하며, \
한 고시 내에 여러 보호구역 유형이 혼재하는 경우가 있어 필지 단위 분류에 활용 가능하다.')
doc.add_paragraph()
table_with_header(doc,
    ['지정유형 (상위 10)', '건수'],
    [(v, c) for v, c in 지정유형_dist]
)

doc.add_paragraph()
heading(doc, '3.3 지방산림청 컬럼', level=2)
para(doc, '2001~2005년대 구명칭(지방산림관리청)과 현행명칭(지방산림청)이 혼재한다. \
또한 148건이 NULL로 추출되어 파일명 기반 보완이 필요하다.')
doc.add_paragraph()
table_with_header(doc,
    ['지방산림청', '건수'],
    [(v if v else '(NULL)', c) for v, c in 기관_dist]
)

doc.add_paragraph()
heading(doc, '3.4 유전자원보호구역_여부 컬럼', level=2)
para(doc, '파일명에 "유전자원" 문자열 포함 여부로 설정된 컬럼이나, 현재 전체 1,442건이 False로 \
기록되어 있다. 파일명 파싱 로직 오류로 확인되며 재처리가 필요하다. \
중장기적으로는 지정유형 컬럼을 기반으로 필지 단위 판별로 전환하는 것이 정확하다.')
doc.add_paragraph()
table_with_header(doc,
    ['값', '건수'],
    [(str(v), c) for v, c in 유전자원여부]
)


# ── 4. 데이터 형식 비일관성 ──
doc.add_page_break()
heading(doc, '4. 데이터 형식 비일관성')

heading(doc, '4.1 고시날짜', level=2)
para(doc, '6가지 날짜 형식이 혼재한다. DATE 타입 통합 또는 표준 형식으로 정규화가 필요하다.')
doc.add_paragraph()
table_with_header(doc,
    ['날짜 형식 패턴', '건수'],
    [(p, c) for p, c in 날짜패턴]
)

doc.add_paragraph()
heading(doc, '4.2 면적 컬럼 (지정면적·해제면적·잔여면적)', level=2)
para(doc, '대부분 단위 없이 순수 숫자로 기록되어 있으며, 일부는 ㎡ 또는 ha를 명시한다. \
단위 불명시 건의 경우 맥락상 ㎡로 추정되나 확인이 필요하다. \
숫자 변환 전 콤마 제거, 단위 분리, ha→㎡ 변환 처리가 필요하다.')
doc.add_paragraph()
table_with_header(doc,
    ['면적값 유형', '건수'],
    [(t, c) for t, c in 면적단위]
)

doc.add_paragraph()
heading(doc, '4.3 소재지', level=2)
para(doc, '광역시도 포함 여부가 일관되지 않는다. 시도명 미포함 건은 고시 발급 지방산림청의 \
관할 지역을 참고하여 보완하거나, 주소 정규화 API를 통해 표준화할 수 있다.')
doc.add_paragraph()
table_with_header(doc,
    ['소재지 패턴', '건수'],
    [(p, c) for p, c in 소재지패턴]
)

doc.add_paragraph()
heading(doc, '4.4 지번', level=2)
para(doc, '"산 217"처럼 산과 숫자 사이 공백이 포함된 케이스가 일부 존재한다. \
지번 파싱 또는 GIS 연계 시 공백 제거 정규화 필요하다.')
doc.add_paragraph()
table_with_header(doc,
    ['지번 패턴', '건수'],
    [(p, c) for p, c in 지번패턴]
)


# ── 5. 전처리 방향 ──
doc.add_page_break()
heading(doc, '5. 전처리 방향 및 우선순위')

items = [
    ('P1', '유전자원보호구역_여부 재처리',
     '파일명 파싱 버그 수정 후 UPDATE 또는 지정유형 기반으로 필지 단위 재판별.\n'
     '예: 지정유형 LIKE \'%유전자원%\' → TRUE'),
    ('P1', '고시날짜 표준화',
     '6가지 혼재 패턴을 파싱하여 DATE 타입 컬럼으로 변환.\n'
     'YYYY년 MM월 DD일 / YYYY-MM-DD / YYYY.MM.DD 등 정규식 처리'),
    ('P1', '유형 컬럼 정규화',
     f'{유형_unique}개 → 지정/해제/정정/변경/공고 5개 범주로 통합.\n'
     '현재 파일명 파싱 기반 분류 컬럼 활용 또는 LLM 재분류'),
    ('P2', '면적 컬럼 숫자 변환',
     '콤마 제거, 단위 분리(㎡/ha), ha→㎡ 변환하여 NUMERIC 컬럼으로 별도 저장.\n'
     '원본 텍스트 컬럼은 유지'),
    ('P2', '소재지 시도 보완',
     '시도 미포함 건(약 48%)에 대해 지방산림청 관할 지역 매핑 또는\n'
     '도로명주소 API로 표준화'),
    ('P2', '지방산림청 구명칭 통합',
     '~관리청 → ~청 으로 통합 정규화 (역사적 사실 보존을 위해 원본 컬럼 유지 권장)'),
    ('P3', '희박 동적 컬럼 처리',
     'null 비율 95% 이상 컬럼(임반·잔여면적·위치 등)의 활용 여부 결정.\n'
     '공통 컬럼 미달 시 별도 부가정보 테이블로 분리 검토'),
    ('P3', '지번 공백 정규화',
     '"산 217" → "산217" 형식으로 통일'),
    ('P3', 'raw_text 활용 계획',
     '평균 3,013자, 최대 186,953자의 원문 보존 중.\n'
     '필지 추출 누락 의심 건(지정면적 null 등) 재파싱 소스로 활용 가능'),
]

table_with_header(doc,
    ['우선순위', '항목', '상세'],
    items
)


# ── 6. raw_text 통계 ──
doc.add_paragraph()
heading(doc, '6. raw_text 보존 현황', level=1)
para(doc, 'pdfplumber로 추출한 원문 텍스트를 전체 1,442건에 대해 보존하고 있다. '
     '향후 파싱 결과 오류 수정, 추가 필드 재추출 등에 재활용 가능하다.')
doc.add_paragraph()
table_with_header(doc,
    ['통계 항목', '값'],
    [
        ['보존 건수', f'{total_rows:,}건 (100%)'],
        ['최소 길이', f'{int(raw_text_stats[0]):,}자'],
        ['최대 길이', f'{int(raw_text_stats[1]):,}자'],
        ['평균 길이', f'{int(raw_text_stats[2]):,}자'],
        ['중앙값 길이', f'{int(raw_text_stats[3]):,}자'],
    ]
)

# 저장
out_path = '/Users/jigo/Workspace/02_Code/01_Python/fgrr-notice-analysis/docs/전처리_품질분석_보고서.docx'
import pathlib
pathlib.Path(out_path).parent.mkdir(exist_ok=True)
doc.save(out_path)
print(f"저장 완료: {out_path}")
